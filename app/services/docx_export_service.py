"""DOCX export helpers for lesson content."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Any, Callable, Dict, List, Optional

from bs4 import BeautifulSoup

from src.observability.logger import get_logger

logger = get_logger(__name__)

_DOCX_IMPORTS: Optional[Dict[str, Any]] = None


def _get_docx_imports() -> Dict[str, Any]:
    global _DOCX_IMPORTS
    if _DOCX_IMPORTS is not None:
        return _DOCX_IMPORTS

    try:
        from docx import Document as DocxDocument
        from docx.enum.text import WD_ALIGN_PARAGRAPH as DocxAlign
        from docx.oxml.ns import qn as docx_qn
        from docx.shared import Inches as DocxInches, Pt as DocxPt, RGBColor as DocxRGBColor
    except ImportError as e:
        raise RuntimeError(
            "DOCX 导出依赖缺失，请安装 python-docx 后重试"
        ) from e

    _DOCX_IMPORTS = {
        "Document": DocxDocument,
        "WD_ALIGN_PARAGRAPH": DocxAlign,
        "qn": docx_qn,
        "Inches": DocxInches,
        "Pt": DocxPt,
        "RGBColor": DocxRGBColor,
    }
    return _DOCX_IMPORTS


def _set_docx_style_font(style: Any, east_asia_font: str, western_font: str, size_pt: int, *, bold: bool = False) -> None:
    imports = _get_docx_imports()
    font = style.font
    font.name = western_font
    font.size = imports["Pt"](size_pt)
    font.bold = bold
    font.color.rgb = imports["RGBColor"](0, 0, 0)
    style.element.rPr.rFonts.set(imports["qn"]("w:eastAsia"), east_asia_font)
    style.element.rPr.rFonts.set(imports["qn"]("w:ascii"), western_font)
    style.element.rPr.rFonts.set(imports["qn"]("w:hAnsi"), western_font)


def _configure_docx_styles(document: Any) -> None:
    body_cjk_font = "SimSun"
    latin_font = "Times New Roman"

    _set_docx_style_font(document.styles["Normal"], body_cjk_font, latin_font, 12)
    _set_docx_style_font(document.styles["List Bullet"], body_cjk_font, latin_font, 12)
    _set_docx_style_font(document.styles["List Number"], body_cjk_font, latin_font, 12)


def _apply_run_font(run: Any, east_asia_font: str, western_font: str, size_pt: int, *, bold: bool = False) -> None:
    imports = _get_docx_imports()
    font = run.font
    font.name = east_asia_font
    font.size = imports["Pt"](size_pt)
    font.bold = bold
    font.color.rgb = imports["RGBColor"](0, 0, 0)

    run_element = getattr(run, "_element", None)
    if run_element is None:
        return
    if hasattr(run_element, "get_or_add_rPr"):
        run_properties = run_element.get_or_add_rPr()
    else:
        run_properties = getattr(run_element, "rPr", None)
    if run_properties is not None and getattr(run_properties, "rFonts", None) is not None:
        run_properties.rFonts.set(imports["qn"]("w:eastAsia"), east_asia_font)
        run_properties.rFonts.set(imports["qn"]("w:ascii"), western_font)
        run_properties.rFonts.set(imports["qn"]("w:hAnsi"), western_font)
        run_properties.rFonts.set(imports["qn"]("w:cs"), western_font)


def _looks_like_heading_1_text(text: str) -> bool:
    clean = " ".join(str(text or "").split())
    if len(clean) > 30:
        return False
    return bool(re.match(r"^[一二三四五六七八九十]+[、\.．]\s*[^。；：:]{1,20}$", clean))


def _apply_paragraph_paper_style(paragraph: Any, role: str) -> None:
    imports = _get_docx_imports()
    format_ = paragraph.paragraph_format
    format_.line_spacing = 1.5

    body_cjk_font = "SimSun"
    heading_cjk_font = "SimHei"
    latin_font = "Times New Roman"

    if role == "title":
        format_.space_before = imports["Pt"](0)
        format_.space_after = imports["Pt"](18)
        paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
        for run in getattr(paragraph, "runs", []):
            _apply_run_font(run, heading_cjk_font, latin_font, 16, bold=True)
        return

    if role == "heading_1":
        format_.space_before = imports["Pt"](12)
        format_.space_after = imports["Pt"](6)
        for run in getattr(paragraph, "runs", []):
            _apply_run_font(run, heading_cjk_font, latin_font, 14, bold=True)
        return

    if role == "heading_2":
        format_.space_before = imports["Pt"](10)
        format_.space_after = imports["Pt"](4)
        for run in getattr(paragraph, "runs", []):
            _apply_run_font(run, body_cjk_font, latin_font, 12, bold=True)
        return

    if role == "heading_3":
        format_.space_before = imports["Pt"](8)
        format_.space_after = imports["Pt"](2)
        for run in getattr(paragraph, "runs", []):
            _apply_run_font(run, body_cjk_font, latin_font, 12, bold=True)
        return

    format_.space_before = imports["Pt"](0)
    format_.space_after = imports["Pt"](6)
    for run in getattr(paragraph, "runs", []):
        _apply_run_font(run, body_cjk_font, latin_font, 12, bold=False)


def _normalize_image_to_png_stream(image: Any) -> Optional[BytesIO]:
    try:
        from PIL import ImageOps
    except ImportError:
        return None

    try:
        normalized = ImageOps.exif_transpose(image)
        normalized.load()
        if normalized.mode not in {"RGB", "L"}:
            normalized = normalized.convert("RGB")
        # Keep exported lesson images lightweight to reduce DOCX memory usage.
        normalized.thumbnail((500, 500))

        buffer = BytesIO()
        normalized.save(buffer, format="PNG")
        buffer.seek(0)
        return buffer
    except Exception as e:
        logger.warning("lesson_docx.normalize_image_failed error=%r", e)
        return None


def _build_docx_compatible_image_stream(image_path: Path) -> Optional[BytesIO]:
    try:
        from PIL import Image, ImageFile
    except ImportError:
        logger.warning("lesson_docx.pillow_missing path=%s", image_path)
        return None

    try:
        # Some textbook-exported images are progressive/CMYK/truncated.
        # Let Pillow decode them as permissively as possible, then normalize.
        ImageFile.LOAD_TRUNCATED_IMAGES = True
        with Image.open(image_path) as image:
            return _normalize_image_to_png_stream(image)
    except Exception as e:
        logger.warning("lesson_docx.convert_image_failed path=%s error=%r", image_path, e)
        return None


def _build_docx_compatible_image_stream_from_bytes(image_bytes: bytes, *, src: str = "") -> Optional[BytesIO]:
    if not image_bytes:
        return None
    try:
        from PIL import Image, ImageFile
    except ImportError:
        logger.warning("lesson_docx.pillow_missing embedded_src=%s", src)
        return None

    try:
        ImageFile.LOAD_TRUNCATED_IMAGES = True
        with Image.open(BytesIO(image_bytes)) as image:
            return _normalize_image_to_png_stream(image)
    except Exception as e:
        logger.warning("lesson_docx.convert_embedded_image_failed src=%s error=%r", src, e)
        return None


def _decode_embedded_image(src: str, resolve_image_bytes: Callable[[str], Optional[bytes]]) -> Optional[BytesIO]:
    image_bytes = resolve_image_bytes(src)
    if not image_bytes:
        return None
    normalized_stream = _build_docx_compatible_image_stream_from_bytes(image_bytes, src=src)
    if normalized_stream is not None:
        return normalized_stream
    try:
        return BytesIO(image_bytes)
    except Exception:
        return None


def _append_text_paragraph(
    document: Any,
    text: str,
    *,
    style: Optional[str] = None,
    role: str = "body",
) -> None:
    clean = " ".join(str(text or "").split())
    if not clean:
        return
    if role == "body" and _looks_like_heading_1_text(clean):
        role = "heading_1"
    paragraph = document.add_paragraph(style=style)
    paragraph.add_run(clean)
    _apply_paragraph_paper_style(paragraph, role)


def _append_image_to_docx(
    document: Any,
    src: str,
    alt_text: str,
    resolve_image_path: Callable[[str], Optional[Path]],
    resolve_image_bytes: Optional[Callable[[str], Optional[bytes]]] = None,
) -> None:
    imports = _get_docx_imports()
    if resolve_image_bytes is not None:
        embedded_stream = _decode_embedded_image(src, resolve_image_bytes)
        if embedded_stream is not None:
            try:
                paragraph = document.add_paragraph()
                paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                run = paragraph.add_run()
                run.add_picture(embedded_stream, width=imports["Inches"](3.8))
                if alt_text:
                    caption = document.add_paragraph()
                    caption.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                    caption.add_run(alt_text)
                return
            except Exception as embedded_error:
                logger.warning("lesson_docx.embedded_image_failed src=%s error=%r", src, embedded_error)

    image_path = resolve_image_path(src)
    if image_path and image_path.exists():
        suffix = image_path.suffix.lower()
        prefer_normalized_insert = suffix in {".jpg", ".jpeg", ".jpx", ".jp2", ".j2k", ".jpf", ".tif", ".tiff"}
        if prefer_normalized_insert:
            converted_stream = _build_docx_compatible_image_stream(image_path)
            if converted_stream is not None:
                try:
                    paragraph = document.add_paragraph()
                    paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                    run = paragraph.add_run()
                    run.add_picture(converted_stream, width=imports["Inches"](3.8))
                    if alt_text:
                        caption = document.add_paragraph()
                        caption.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                        caption.add_run(alt_text)
                    return
                except Exception as normalized_error:
                    logger.warning(
                        "lesson_docx.normalized_image_failed path=%s error=%r",
                        image_path,
                        normalized_error,
                    )

        try:
            paragraph = document.add_paragraph()
            paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=imports["Inches"](3.8))
            if alt_text:
                caption = document.add_paragraph()
                caption.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                caption.add_run(alt_text)
        except Exception as e:
            logger.warning("lesson_docx.direct_image_failed path=%s error=%r", image_path, e)
            converted_stream = _build_docx_compatible_image_stream(image_path)
            if converted_stream is not None:
                try:
                    paragraph = document.add_paragraph()
                    paragraph.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                    run = paragraph.add_run()
                    run.add_picture(converted_stream, width=imports["Inches"](3.8))
                    if alt_text:
                        caption = document.add_paragraph()
                        caption.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                        caption.add_run(alt_text)
                    return
                except Exception as converted_error:
                    logger.warning(
                        "lesson_docx.converted_image_failed path=%s error=%r",
                        image_path,
                        converted_error,
                    )

            if alt_text:
                fallback = document.add_paragraph()
                fallback.alignment = imports["WD_ALIGN_PARAGRAPH"].CENTER
                fallback.add_run(f"[图片未导出] {alt_text}")


def _append_node_content(
    document: Any,
    node: Any,
    resolve_image_path: Callable[[str], Optional[Path]],
    resolve_image_bytes: Optional[Callable[[str], Optional[bytes]]] = None,
    *,
    style: Optional[str] = None,
    role: str = "body",
) -> None:
    text_buffer: List[str] = []

    for child in getattr(node, "children", []):
        child_name = getattr(child, "name", None)
        if child_name == "img":
            _append_text_paragraph(document, " ".join(text_buffer), style=style, role=role)
            text_buffer = []
            _append_image_to_docx(
                document,
                child.get("src", ""),
                child.get("alt", "").strip(),
                resolve_image_path,
                resolve_image_bytes,
            )
            continue

        if isinstance(child, str):
            text_buffer.append(child)
            continue

        child_text = child.get_text(" ", strip=True)
        if child_text:
            text_buffer.append(child_text)

    _append_text_paragraph(document, " ".join(text_buffer), style=style, role=role)


def _render_html_to_docx(
    document: Any,
    html: str,
    resolve_image_path: Callable[[str], Optional[Path]],
    resolve_image_bytes: Optional[Callable[[str], Optional[bytes]]] = None,
) -> None:
    soup = BeautifulSoup(html or "", "html.parser")
    root_nodes = soup.contents if soup.contents else []

    for node in root_nodes:
        if isinstance(node, str):
            _append_text_paragraph(document, node)
            continue

        name = getattr(node, "name", "") or ""

        if name == "h1":
            _append_node_content(document, node, resolve_image_path, resolve_image_bytes, role="title")
            continue
        if name == "h2":
            _append_node_content(document, node, resolve_image_path, resolve_image_bytes, role="heading_1")
            continue
        if name == "h3":
            _append_node_content(document, node, resolve_image_path, resolve_image_bytes, role="heading_2")
            continue
        if name == "h4":
            _append_node_content(document, node, resolve_image_path, resolve_image_bytes, role="heading_3")
            continue
        if name in {"p", "blockquote", "div"}:
            _append_node_content(document, node, resolve_image_path, resolve_image_bytes)
            continue
        if name == "hr":
            document.add_paragraph("")
            continue
        if name in {"ul", "ol"}:
            list_style = "List Bullet" if name == "ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                _append_node_content(document, li, resolve_image_path, resolve_image_bytes, style=list_style)
            continue
        if name == "img":
            _append_image_to_docx(
                document,
                node.get("src", ""),
                node.get("alt", "").strip(),
                resolve_image_path,
                resolve_image_bytes,
            )
            continue

        _append_node_content(document, node, resolve_image_path, resolve_image_bytes)


def build_lesson_docx_bytes(
    *,
    content_html: str,
    resolve_image_path: Callable[[str], Optional[Path]],
    resolve_image_bytes: Optional[Callable[[str], Optional[bytes]]] = None,
) -> bytes:
    imports = _get_docx_imports()
    document = imports["Document"]()
    section = document.sections[0]
    section.top_margin = imports["Inches"](0.8)
    section.bottom_margin = imports["Inches"](0.8)
    section.left_margin = imports["Inches"](0.9)
    section.right_margin = imports["Inches"](0.9)
    _configure_docx_styles(document)

    _render_html_to_docx(document, content_html, resolve_image_path, resolve_image_bytes)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
